"""
Script pour générer la documentation Word du projet OCP
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Créer le document
doc = Document()

# ============================================
# STYLES
# ============================================
# Style pour les titres
style = doc.styles['Heading 1']
style.font.size = Pt(18)
style.font.bold = True
style.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # Vert OCP

style2 = doc.styles['Heading 2']
style2.font.size = Pt(14)
style2.font.bold = True
style2.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

style3 = doc.styles['Heading 3']
style3.font.size = Pt(12)
style3.font.bold = True
style3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================
# PAGE DE TITRE
# ============================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("\n\n\n")
run = title.add_run("📊 DOCUMENTATION TECHNIQUE")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Application de Gestion de Bases de Données")
run.font.size = Pt(20)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run("Port de Jorf Lasfar - OCP")
run.font.size = Pt(16)
run.font.italic = True

doc.add_paragraph("\n\n\n")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Version 1.0\n").font.size = Pt(12)
info.add_run("Décembre 2025\n").font.size = Pt(12)

doc.add_page_break()

# ============================================
# TABLE DES MATIÈRES
# ============================================
doc.add_heading("Table des Matières", level=1)
toc_items = [
    "1. Introduction",
    "2. Architecture du Projet",
    "3. Structure des Fichiers",
    "   3.1 Fichiers Racine",
    "   3.2 Backend (Django)",
    "   3.3 Frontend (React)",
    "4. Fonctionnement de l'Application",
    "   4.1 Authentification",
    "   4.2 Gestion des Fichiers Excel",
    "   4.3 Système de Cache",
    "   4.4 API REST",
    "5. Guide d'Utilisation",
    "6. Endpoints API",
    "7. Base de Données",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.5) if item.startswith("   ") else Inches(0)

doc.add_page_break()

# ============================================
# 1. INTRODUCTION
# ============================================
doc.add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "Cette application est un système fullstack de gestion de bases de données Excel "
    "développé pour l'OCP (Office Chérifien des Phosphates), spécifiquement pour gérer "
    "les données Import et Export du Port de Jorf Lasfar."
)

doc.add_heading("Objectifs de l'application", level=2)
objectives = [
    "Centraliser les données d'import/export dans une interface web moderne",
    "Permettre la consultation et modification des fichiers Excel en temps réel",
    "Assurer la traçabilité des modifications (qui a modifié quoi et quand)",
    "Offrir un système d'archivage sécurisé avec possibilité de restauration",
    "Gérer les accès utilisateurs avec différents niveaux de permissions"
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("Technologies utilisées", level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
headers = table.rows[0].cells
headers[0].text = "Composant"
headers[1].text = "Technologies"
data = [
    ("Backend", "Django 5.2, Django REST Framework, SQLite"),
    ("Frontend", "React 18, Vite, React Router"),
    ("Authentification", "JWT (JSON Web Tokens)"),
    ("Manipulation Excel", "openpyxl"),
]
for i, (comp, tech) in enumerate(data, 1):
    table.rows[i].cells[0].text = comp
    table.rows[i].cells[1].text = tech

doc.add_page_break()

# ============================================
# 2. ARCHITECTURE
# ============================================
doc.add_heading("2. Architecture du Projet", level=1)

doc.add_paragraph(
    "L'application suit une architecture client-serveur classique avec séparation "
    "claire entre le frontend (interface utilisateur) et le backend (API et logique métier)."
)

doc.add_heading("Schéma d'architecture", level=2)
architecture = """
┌─────────────────────────────────────────────────────────────┐
│                      UTILISATEUR                             │
│                    (Navigateur Web)                          │
└─────────────────────────┬───────────────────────────────────┘
                          │ HTTP/HTTPS
                          ▼
┌─────────────────────────────────────────────────────────────┐
│                   FRONTEND (React + Vite)                    │
│                   Port: 5173 (dev)                           │
│  ┌─────────────────────────────────────────────────────┐    │
│  │  Pages: Login, Home, FileSheets, SheetDetail, Users │    │
│  │  Services: api.js (client HTTP)                      │    │
│  └─────────────────────────────────────────────────────┘    │
└─────────────────────────┬───────────────────────────────────┘
                          │ API REST (JSON)
                          ▼
┌─────────────────────────────────────────────────────────────┐
│                   BACKEND (Django REST)                      │
│                   Port: 8000                                 │
│  ┌─────────────────────────────────────────────────────┐    │
│  │  views.py: Logique CRUD fichiers Excel              │    │
│  │  views_users.py: Gestion utilisateurs               │    │
│  │  models.py: FileCache, SheetDataCache               │    │
│  └─────────────────────────────────────────────────────┘    │
└──────────────┬─────────────────────────┬────────────────────┘
               │                         │
               ▼                         ▼
┌──────────────────────────┐  ┌──────────────────────────────┐
│   Base de données        │  │    Fichiers Excel (.xlsx)    │
│   SQLite (db.sqlite3)    │  │    Dossier racine du projet  │
│   - Utilisateurs         │  │    - Import 2025.xlsx        │
│   - Cache fichiers       │  │    - Export 2025.xlsx        │
│   - Cache feuilles       │  │    - _archives/              │
└──────────────────────────┘  └──────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(architecture)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ============================================
# 3. STRUCTURE DES FICHIERS
# ============================================
doc.add_heading("3. Structure des Fichiers", level=1)

# 3.1 Fichiers Racine
doc.add_heading("3.1 Fichiers Racine", level=2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier/Dossier"
table.rows[0].cells[1].text = "Type"
table.rows[0].cells[2].text = "Description"
root_files = [
    ("_archives/", "Dossier", "Stockage des fichiers Excel archivés (supprimés mais récupérables)"),
    ("backend/", "Dossier", "Code source du serveur Django (API REST)"),
    ("frontend/", "Dossier", "Code source de l'interface React"),
    ("*.xlsx", "Fichiers", "Fichiers Excel de données Import/Export"),
    ("requirements.txt", "Config", "Liste des dépendances Python (pip install -r requirements.txt)"),
    ("create_shortcut.ps1", "Script", "Script PowerShell pour créer un raccourci Windows"),
    ("update_icon.ps1", "Script", "Script PowerShell pour mettre à jour l'icône"),
    ("SCRIPT_PRESENTATION.md", "Doc", "Script de présentation du projet"),
]
for i, (file, type_, desc) in enumerate(root_files, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = type_
    table.rows[i].cells[2].text = desc

# 3.2 Backend
doc.add_heading("3.2 Backend (Django)", level=2)

doc.add_heading("Configuration Django (backend/backend/)", level=3)
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier"
table.rows[0].cells[1].text = "Rôle"
backend_config = [
    ("__init__.py", "Marque le dossier comme module Python"),
    ("settings.py", "Configuration principale: BDD, CORS, JWT, apps installées, middleware"),
    ("urls.py", "Routes principales de l'application (/api/, /admin/)"),
    ("asgi.py", "Configuration pour serveur ASGI (communications asynchrones)"),
    ("wsgi.py", "Configuration pour serveur WSGI (production)"),
]
for i, (file, role) in enumerate(backend_config, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = role

doc.add_heading("Application API (backend/api/)", level=3)
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier"
table.rows[0].cells[1].text = "Rôle"
api_files = [
    ("models.py", "Définition des modèles de données:\n- FileCache: métadonnées des fichiers Excel\n- SheetDataCache: données des feuilles en cache\n- ExcelFile, ExcelColumn: structure des fichiers"),
    ("views.py", "Logique métier principale:\n- CRUD fichiers Excel\n- Lecture/écriture des feuilles\n- Système de cache\n- Import/Export fichiers"),
    ("views_users.py", "Gestion des utilisateurs:\n- Création, modification, suppression\n- Changement de mot de passe\n- Authentification JWT"),
    ("serializers.py", "Conversion des modèles en JSON et vice-versa pour l'API REST"),
    ("urls.py", "Définition des routes API (/files/, /sheets/, /users/, etc.)"),
    ("admin.py", "Configuration de l'interface d'administration Django"),
    ("apps.py", "Configuration de l'application Django"),
    ("tests.py", "Tests unitaires (à développer)"),
]
for i, (file, role) in enumerate(api_files, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = role

doc.add_heading("Scripts utilitaires (backend/)", level=3)
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier"
table.rows[0].cells[1].text = "Rôle"
scripts = [
    ("manage.py", "Point d'entrée Django: runserver, migrate, createsuperuser, shell"),
    ("db.sqlite3", "Base de données SQLite contenant utilisateurs et cache"),
    ("create_user.py", "Script pour créer un utilisateur en ligne de commande"),
    ("check_files.py", "Script de vérification des fichiers Excel"),
    ("refresh_column_types.py", "Script pour rafraîchir les types de colonnes dans le cache"),
]
for i, (file, role) in enumerate(scripts, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = role

doc.add_page_break()

# 3.3 Frontend
doc.add_heading("3.3 Frontend (React)", level=2)

doc.add_heading("Configuration (frontend/)", level=3)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier"
table.rows[0].cells[1].text = "Rôle"
frontend_config = [
    ("index.html", "Page HTML principale, point d'entrée de l'application"),
    ("package.json", "Dépendances npm et scripts (npm run dev, npm run build)"),
    ("package-lock.json", "Versions exactes des dépendances pour reproductibilité"),
    ("vite.config.js", "Configuration de Vite: proxy API, port de développement"),
    ("eslint.config.js", "Règles de qualité et style de code JavaScript"),
    ("README.md", "Documentation du frontend"),
]
for i, (file, role) in enumerate(frontend_config, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = role

doc.add_heading("Code source (frontend/src/)", level=3)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Fichier"
table.rows[0].cells[1].text = "Rôle"
src_files = [
    ("main.jsx", "Point d'entrée React: monte l'application et configure le Router"),
    ("App.jsx", "Définition des routes: quelle page afficher selon l'URL"),
    ("App.css", "Styles globaux de l'application"),
    ("index.css", "Styles de base (reset CSS, polices, variables)"),
]
for i, (file, role) in enumerate(src_files, 1):
    table.rows[i].cells[0].text = file
    table.rows[i].cells[1].text = role

doc.add_heading("Pages (frontend/src/pages/)", level=3)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Page"
table.rows[0].cells[1].text = "Fonctionnalité"
pages = [
    ("login.jsx", "Page de connexion avec authentification JWT"),
    ("Home.jsx", "Page d'accueil: liste des fichiers, import, création, archives"),
    ("FileSheets.jsx", "Liste des feuilles d'un fichier Excel sélectionné"),
    ("SheetDetail.jsx", "Tableau de données avec CRUD (ajout, modification, suppression)"),
    ("CreateFile.jsx", "Formulaire de création de nouveau fichier avec colonnes personnalisées"),
    ("Users.jsx", "Gestion des utilisateurs (réservé aux administrateurs)"),
    ("*.css", "Styles spécifiques à chaque page"),
]
for i, (page, func) in enumerate(pages, 1):
    table.rows[i].cells[0].text = page
    table.rows[i].cells[1].text = func

doc.add_heading("Services (frontend/src/services/)", level=3)
doc.add_paragraph(
    "api.js - Client HTTP centralisé pour toutes les communications avec le backend:"
)
services = [
    "authService: login(), logout(), refreshToken()",
    "userService: getCurrentUser(), getUsers(), createUser(), updateUser()",
    "filesService: getFiles(), importFile(), createFile(), deleteFile()",
    "sheetsService: getSheets(), getData(), addEntry(), updateEntry(), deleteEntry()"
]
for service in services:
    doc.add_paragraph(service, style='List Bullet')

doc.add_page_break()

# ============================================
# 4. FONCTIONNEMENT
# ============================================
doc.add_heading("4. Fonctionnement de l'Application", level=1)

# 4.1 Authentification
doc.add_heading("4.1 Authentification", level=2)
doc.add_paragraph(
    "L'application utilise JWT (JSON Web Tokens) pour sécuriser les accès:"
)
auth_steps = [
    "1. L'utilisateur entre son username et mot de passe sur la page de connexion",
    "2. Le frontend envoie une requête POST à /api/login/",
    "3. Le backend vérifie les credentials et génère deux tokens:",
    "   - Access Token (durée courte ~15min): pour authentifier les requêtes",
    "   - Refresh Token (durée longue ~24h): pour renouveler l'access token",
    "4. Les tokens sont stockés dans localStorage",
    "5. Chaque requête API inclut l'Access Token dans le header Authorization",
    "6. Quand l'Access Token expire, le Refresh Token est utilisé pour en obtenir un nouveau"
]
for step in auth_steps:
    doc.add_paragraph(step)

doc.add_heading("Types d'utilisateurs", level=3)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Rôle"
table.rows[0].cells[1].text = "Permissions"
table.rows[1].cells[0].text = "Utilisateur"
table.rows[1].cells[1].text = "Consulter, ajouter, modifier des données"
table.rows[2].cells[0].text = "Administrateur"
table.rows[2].cells[1].text = "Tout + gérer utilisateurs + suppression définitive"

# 4.2 Gestion des fichiers
doc.add_heading("4.2 Gestion des Fichiers Excel", level=2)

doc.add_heading("Import d'un fichier existant", level=3)
import_steps = [
    "1. L'utilisateur clique sur 'Importer' et sélectionne un fichier .xlsx",
    "2. Le fichier est envoyé au backend via multipart/form-data",
    "3. Le backend valide que c'est un fichier Excel valide avec openpyxl",
    "4. Le fichier est sauvegardé dans le dossier racine du projet",
    "5. Les métadonnées sont extraites et stockées dans FileCache",
    "6. Les données de chaque feuille sont mises en cache dans SheetDataCache",
    "7. Le fichier apparaît dans la liste de la page d'accueil"
]
for step in import_steps:
    doc.add_paragraph(step)

doc.add_heading("Création d'un nouveau fichier", level=3)
create_steps = [
    "1. L'utilisateur clique sur 'Nouveau fichier'",
    "2. Il définit le nom du fichier et les colonnes (nom, type, obligatoire)",
    "3. Le backend crée un fichier Excel avec openpyxl",
    "4. Les en-têtes sont formatés avec style (couleur verte, gras)",
    "5. Le fichier est sauvegardé et mis en cache"
]
for step in create_steps:
    doc.add_paragraph(step)

doc.add_heading("Archivage (Soft Delete)", level=3)
doc.add_paragraph(
    "Quand un fichier est 'supprimé', il n'est pas réellement effacé:"
)
archive_steps = [
    "Le fichier physique est déplacé vers le dossier _archives/",
    "L'entrée dans FileCache est marquée is_deleted=True",
    "La date et l'utilisateur qui a supprimé sont enregistrés",
    "Le fichier peut être restauré depuis l'interface Archives",
    "Les données en cache sont conservées pour historique"
]
for step in archive_steps:
    doc.add_paragraph(step, style='List Bullet')

# 4.3 Système de Cache
doc.add_heading("4.3 Système de Cache", level=2)
doc.add_paragraph(
    "Pour optimiser les performances, l'application utilise un système de cache à deux niveaux:"
)

doc.add_heading("FileCache - Métadonnées des fichiers", level=3)
filecache_fields = [
    "filename: Nom du fichier",
    "name: Nom affiché (sans extension)",
    "file_path: Chemin complet sur le disque",
    "sheets_count: Nombre de feuilles",
    "sheets_json: Liste des noms de feuilles",
    "sheets_details: Détails (colonnes, entrées) par feuille",
    "total_entries: Nombre total d'entrées",
    "file_size: Taille en octets",
    "file_modified: Date de dernière modification",
    "last_modified_by: Dernier utilisateur qui a modifié"
]
for field in filecache_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading("SheetDataCache - Données des feuilles", level=3)
sheetcache_fields = [
    "file_cache: Référence au fichier parent",
    "sheet_name: Nom de la feuille",
    "headers: Liste des en-têtes de colonnes",
    "columns_info: Types et propriétés des colonnes",
    "data: Données JSON de toutes les lignes",
    "rows_count: Nombre de lignes"
]
for field in sheetcache_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading("Synchronisation du cache", level=3)
doc.add_paragraph(
    "À chaque accès à la liste des fichiers, le système vérifie si les fichiers "
    "sur le disque ont été modifiés depuis le dernier cache. Si oui, le cache est "
    "automatiquement mis à jour. Cela garantit que les données affichées sont "
    "toujours synchronisées avec les fichiers réels."
)

doc.add_page_break()

# 4.4 API REST
doc.add_heading("4.4 API REST", level=2)
doc.add_paragraph(
    "Le backend expose une API REST complète pour toutes les opérations. "
    "Toutes les routes (sauf /login/) nécessitent un token JWT valide."
)

# ============================================
# 5. GUIDE D'UTILISATION
# ============================================
doc.add_heading("5. Guide d'Utilisation", level=1)

doc.add_heading("Démarrage de l'application", level=2)
startup_steps = [
    "1. Ouvrir un terminal dans le dossier du projet",
    "2. Démarrer le backend: cd backend && python manage.py runserver",
    "3. Ouvrir un second terminal",
    "4. Démarrer le frontend: cd frontend && npm run dev",
    "5. Ouvrir http://localhost:5173 dans le navigateur",
    "6. Se connecter avec ses identifiants"
]
for step in startup_steps:
    doc.add_paragraph(step)

doc.add_heading("Flux de travail typique", level=2)
workflow = [
    "1. CONNEXION: Entrer username et mot de passe",
    "2. PAGE D'ACCUEIL: Visualiser les fichiers disponibles",
    "3. SÉLECTION: Cliquer sur un fichier pour voir ses feuilles",
    "4. FEUILLE: Cliquer sur une feuille pour voir les données",
    "5. AJOUT: Cliquer sur 'Ajouter une entrée' pour créer une ligne",
    "6. MODIFICATION: Cliquer sur une ligne pour la modifier",
    "7. SUPPRESSION: Utiliser l'icône corbeille pour supprimer"
]
for step in workflow:
    doc.add_paragraph(step)

doc.add_page_break()

# ============================================
# 6. ENDPOINTS API
# ============================================
doc.add_heading("6. Endpoints API", level=1)

doc.add_heading("Authentification", level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Méthode"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Description"
auth_endpoints = [
    ("POST", "/api/login/", "Connexion, retourne les tokens JWT"),
    ("POST", "/api/token/refresh/", "Rafraîchir l'access token"),
    ("GET", "/api/me/", "Informations de l'utilisateur connecté"),
]
for i, (method, endpoint, desc) in enumerate(auth_endpoints, 1):
    table.rows[i].cells[0].text = method
    table.rows[i].cells[1].text = endpoint
    table.rows[i].cells[2].text = desc

doc.add_heading("Gestion des fichiers", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Méthode"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Description"
file_endpoints = [
    ("GET", "/api/files/", "Liste tous les fichiers Excel"),
    ("POST", "/api/files/create/", "Créer un nouveau fichier"),
    ("POST", "/api/files/import/", "Importer un fichier existant"),
    ("POST", "/api/files/refresh/", "Rafraîchir le cache"),
    ("DELETE", "/api/files/{filename}/delete/", "Archiver un fichier"),
    ("GET", "/api/files/{filename}/download/", "Télécharger un fichier"),
]
for i, (method, endpoint, desc) in enumerate(file_endpoints, 1):
    table.rows[i].cells[0].text = method
    table.rows[i].cells[1].text = endpoint
    table.rows[i].cells[2].text = desc

doc.add_heading("Gestion des feuilles", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Méthode"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Description"
sheet_endpoints = [
    ("GET", "/api/files/{f}/sheets/", "Liste des feuilles d'un fichier"),
    ("POST", "/api/files/{f}/sheets/create/", "Créer une nouvelle feuille"),
    ("GET", "/api/files/{f}/sheets/{s}/columns/", "Colonnes d'une feuille"),
    ("GET", "/api/files/{f}/sheets/{s}/data/", "Données d'une feuille"),
    ("POST", "/api/files/{f}/sheets/{s}/add/", "Ajouter une entrée"),
    ("PUT", "/api/files/{f}/sheets/{s}/update/", "Modifier une entrée"),
]
for i, (method, endpoint, desc) in enumerate(sheet_endpoints, 1):
    table.rows[i].cells[0].text = method
    table.rows[i].cells[1].text = endpoint
    table.rows[i].cells[2].text = desc

doc.add_heading("Gestion des utilisateurs", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Méthode"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Description"
user_endpoints = [
    ("GET", "/api/users/", "Liste tous les utilisateurs"),
    ("POST", "/api/users/create/", "Créer un utilisateur"),
    ("PUT", "/api/users/{id}/update/", "Modifier un utilisateur"),
    ("DELETE", "/api/users/{id}/delete/", "Supprimer un utilisateur"),
    ("POST", "/api/users/change-password/", "Changer le mot de passe"),
]
for i, (method, endpoint, desc) in enumerate(user_endpoints, 1):
    table.rows[i].cells[0].text = method
    table.rows[i].cells[1].text = endpoint
    table.rows[i].cells[2].text = desc

doc.add_page_break()

# ============================================
# 7. BASE DE DONNÉES
# ============================================
doc.add_heading("7. Base de Données", level=1)

doc.add_paragraph(
    "L'application utilise SQLite comme base de données. Le fichier db.sqlite3 "
    "contient toutes les tables nécessaires au fonctionnement de l'application."
)

doc.add_heading("Schéma des tables", level=2)

doc.add_heading("Table: auth_user (Django)", level=3)
user_fields = [
    "id: Identifiant unique",
    "username: Nom d'utilisateur",
    "password: Mot de passe hashé",
    "email: Adresse email",
    "first_name, last_name: Nom complet",
    "is_staff: Est administrateur",
    "is_active: Compte actif",
    "last_login: Dernière connexion"
]
for field in user_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading("Table: api_filecache", level=3)
filecache_db = [
    "id: Identifiant unique",
    "filename: Nom du fichier (unique)",
    "name: Nom affiché",
    "file_path: Chemin complet",
    "sheets_count: Nombre de feuilles",
    "sheets_json: Liste JSON des feuilles",
    "sheets_details: Détails JSON par feuille",
    "total_entries: Nombre d'entrées",
    "file_size: Taille en octets",
    "file_modified: Date modification fichier",
    "cached_at: Date mise en cache",
    "last_modified_by_id: FK vers user",
    "is_deleted: Supprimé (soft delete)",
    "deleted_at: Date suppression",
    "deleted_by_id: FK vers user",
    "archived_path: Chemin archive"
]
for field in filecache_db:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading("Table: api_sheetdatacache", level=3)
sheetcache_db = [
    "id: Identifiant unique",
    "file_cache_id: FK vers FileCache",
    "sheet_name: Nom de la feuille",
    "headers: JSON des en-têtes",
    "columns_info: JSON infos colonnes",
    "data: JSON des données",
    "rows_count: Nombre de lignes",
    "cached_at: Date mise en cache"
]
for field in sheetcache_db:
    doc.add_paragraph(field, style='List Bullet')

# ============================================
# FIN DU DOCUMENT
# ============================================
doc.add_page_break()
doc.add_heading("Contact et Support", level=1)
doc.add_paragraph("Pour toute question ou assistance technique, veuillez contacter l'équipe de développement.")
doc.add_paragraph("")
doc.add_paragraph("Document généré automatiquement - Décembre 2025")
doc.add_paragraph("OCP - Port de Jorf Lasfar")

# Sauvegarder le document
output_path = os.path.join(os.path.dirname(__file__), "Documentation_Technique_OCP.docx")
doc.save(output_path)
print(f"Document Word cree: {output_path}")

# Ouvrir le document
os.startfile(output_path)
print("Document ouvert dans Word!")

